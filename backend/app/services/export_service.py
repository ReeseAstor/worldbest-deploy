"""Export service – generates DOCX and EPUB files from project content."""

from __future__ import annotations

import io
import logging
import uuid

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.project import Book, Chapter, Project

logger = logging.getLogger("ember.export_service")


class ExportService:
    """Build downloadable manuscript files in DOCX or EPUB format."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    async def export_docx(self, project_id: uuid.UUID, options: dict) -> bytes:
        """Fetch all chapters and build a DOCX manuscript.

        Returns the file content as raw bytes.
        """
        from docx import Document
        from docx.shared import Pt, Inches

        project, chapters = await self._load_project_chapters(project_id)

        doc = Document()

        # Title page
        title_para = doc.add_heading(project.title, level=0)
        if project.synopsis:
            doc.add_paragraph(project.synopsis)
        doc.add_page_break()

        # Chapters
        for chapter in chapters:
            doc.add_heading(
                f"Chapter {chapter.number}: {chapter.title}", level=1
            )
            if chapter.content_text:
                for paragraph_text in chapter.content_text.split("\n\n"):
                    paragraph_text = paragraph_text.strip()
                    if paragraph_text:
                        p = doc.add_paragraph(paragraph_text)
                        p.style.font.size = Pt(12)
            doc.add_page_break()

        # Serialize to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    async def export_epub(self, project_id: uuid.UUID, options: dict) -> bytes:
        """Fetch all chapters and build an EPUB e-book.

        Returns the file content as raw bytes.
        """
        from ebooklib import epub

        project, chapters = await self._load_project_chapters(project_id)

        book = epub.EpubBook()
        book.set_identifier(str(project.id))
        book.set_title(project.title)
        book.set_language("en")

        # Table of contents + spine
        toc: list[epub.EpubHtml] = []
        spine: list[str | epub.EpubHtml] = ["nav"]

        for chapter in chapters:
            ch = epub.EpubHtml(
                title=f"Chapter {chapter.number}: {chapter.title}",
                file_name=f"chapter_{chapter.number}.xhtml",
                lang="en",
            )

            # Build HTML content
            html_parts = [f"<h1>Chapter {chapter.number}: {chapter.title}</h1>"]
            if chapter.content_text:
                for para in chapter.content_text.split("\n\n"):
                    para = para.strip()
                    if para:
                        html_parts.append(f"<p>{para}</p>")

            ch.content = "\n".join(html_parts)
            book.add_item(ch)
            toc.append(ch)
            spine.append(ch)

        book.toc = toc
        book.spine = spine

        # Required EPUB navigation
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Serialize to bytes
        buffer = io.BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)
        return buffer.read()

    # ── Helpers ───────────────────────────────────────────────────────

    async def _load_project_chapters(
        self, project_id: uuid.UUID
    ) -> tuple[Project, list[Chapter]]:
        """Load the project and all its chapters in order."""
        result = await self.db.execute(
            select(Project).where(Project.id == project_id)
        )
        project = result.scalar_one_or_none()
        if project is None:
            raise ValueError(f"Project {project_id} not found")

        # Get all books for this project
        books_result = await self.db.execute(
            select(Book)
            .where(Book.project_id == project_id)
            .order_by(Book.order)
        )
        books = books_result.scalars().all()
        book_ids = [b.id for b in books]

        if not book_ids:
            return project, []

        # Get all chapters across all books
        chapters_result = await self.db.execute(
            select(Chapter)
            .where(Chapter.book_id.in_(book_ids))
            .order_by(Chapter.number)
        )
        chapters = list(chapters_result.scalars().all())

        return project, chapters
